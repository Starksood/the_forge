"""Unit tests for DocumentParser class."""
import pytest
import io
from backend.document_parser import (
    DocumentParser,
    UnsupportedFormatError,
    CorruptedFileError,
)


@pytest.fixture
def parser():
    """Create a DocumentParser instance for testing."""
    return DocumentParser()


class TestDocumentParserBasics:
    """Test basic DocumentParser functionality."""
    
    @pytest.mark.asyncio
    async def test_empty_filename_raises_error(self, parser):
        """Test that empty filename raises ValueError."""
        with pytest.raises(ValueError, match="Filename and content are required"):
            await parser.extract_text("", b"content")
    
    @pytest.mark.asyncio
    async def test_empty_content_raises_error(self, parser):
        """Test that empty content raises ValueError."""
        with pytest.raises(ValueError, match="Filename and content are required"):
            await parser.extract_text("file.txt", b"")
    
    @pytest.mark.asyncio
    async def test_no_extension_raises_error(self, parser):
        """Test that file without extension raises UnsupportedFormatError."""
        with pytest.raises(UnsupportedFormatError, match="File has no extension"):
            await parser.extract_text("noextension", b"content")
    
    @pytest.mark.asyncio
    async def test_unsupported_format_raises_error(self, parser):
        """Test that unsupported format raises UnsupportedFormatError."""
        with pytest.raises(UnsupportedFormatError, match="Unsupported format '.xyz'"):
            await parser.extract_text("file.xyz", b"content")
    
    def test_supported_formats(self, parser):
        """Test that parser declares supported formats."""
        assert "pdf" in parser.SUPPORTED_FORMATS
        assert "docx" in parser.SUPPORTED_FORMATS
        assert "txt" in parser.SUPPORTED_FORMATS


class TestTextFileParsing:
    """Test plain text file parsing."""
    
    @pytest.mark.asyncio
    async def test_parse_simple_text(self, parser):
        """Test parsing simple UTF-8 text file."""
        content = b"Hello, world!\nThis is a test document."
        text, pages = await parser.extract_text("test.txt", content)
        
        assert text == "Hello, world!\nThis is a test document."
        assert pages == 1
    
    @pytest.mark.asyncio
    async def test_parse_large_text(self, parser):
        """Test page estimation for large text files."""
        # Create text larger than CHARS_PER_PAGE (3000)
        content = b"A" * 7000
        text, pages = await parser.extract_text("large.txt", content)
        
        assert len(text) == 7000
        assert pages == 2  # 7000 / 3000 = 2.33, rounded down to 2
    
    @pytest.mark.asyncio
    async def test_parse_empty_text_raises_error(self, parser):
        """Test that empty text file raises CorruptedFileError."""
        with pytest.raises(CorruptedFileError, match="Text file is empty"):
            await parser.extract_text("empty.txt", b"   \n\n  ")
    
    @pytest.mark.asyncio
    async def test_parse_non_utf8_text(self, parser):
        """Test parsing text with non-UTF-8 encoding."""
        # Latin-1 encoded text
        content = "Café résumé".encode("latin-1")
        text, pages = await parser.extract_text("latin.txt", content)
        
        # Should decode successfully (possibly with replacement chars)
        assert text is not None
        assert pages >= 1


class TestPDFParsing:
    """Test PDF file parsing."""
    
    @pytest.mark.asyncio
    async def test_parse_valid_pdf(self, parser):
        """Test parsing a valid PDF file."""
        # Create a minimal valid PDF
        pdf_content = self._create_minimal_pdf()
        
        text, pages = await parser.extract_text("test.pdf", pdf_content)
        
        assert isinstance(text, str)
        assert pages >= 1
    
    @pytest.mark.asyncio
    async def test_parse_corrupted_pdf(self, parser):
        """Test that corrupted PDF raises CorruptedFileError."""
        corrupted_pdf = b"Not a real PDF file"
        
        with pytest.raises(CorruptedFileError, match="PDF"):
            await parser.extract_text("corrupted.pdf", corrupted_pdf)
    
    @pytest.mark.asyncio
    async def test_parse_empty_pdf_raises_error(self, parser):
        """Test that PDF with no pages raises CorruptedFileError."""
        # This would need a specially crafted PDF with 0 pages
        # For now, we'll test with invalid PDF that triggers the error path
        with pytest.raises(CorruptedFileError):
            await parser.extract_text("empty.pdf", b"%PDF-invalid")
    
    def _create_minimal_pdf(self):
        """Create a minimal valid PDF for testing."""
        try:
            import PyPDF2
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
        except ImportError:
            pytest.skip("PyPDF2 or reportlab not available")
        
        # Create PDF in memory
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 730, "This is a test page.")
        c.showPage()
        c.save()
        
        return buffer.getvalue()


class TestDOCXParsing:
    """Test DOCX file parsing."""
    
    @pytest.mark.asyncio
    async def test_parse_valid_docx(self, parser):
        """Test parsing a valid DOCX file."""
        docx_content = self._create_minimal_docx()
        
        text, pages = await parser.extract_text("test.docx", docx_content)
        
        assert isinstance(text, str)
        assert len(text) > 0
        assert pages >= 1
    
    @pytest.mark.asyncio
    async def test_parse_corrupted_docx(self, parser):
        """Test that corrupted DOCX raises CorruptedFileError."""
        corrupted_docx = b"Not a real DOCX file"
        
        with pytest.raises(CorruptedFileError, match="DOCX"):
            await parser.extract_text("corrupted.docx", corrupted_docx)
    
    def _create_minimal_docx(self):
        """Create a minimal valid DOCX for testing."""
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not available")
        
        # Create DOCX in memory
        doc = docx.Document()
        doc.add_paragraph("Test DOCX Document")
        doc.add_paragraph("This is a test paragraph.")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()


class TestErrorHandling:
    """Test error handling for various edge cases."""
    
    @pytest.mark.asyncio
    async def test_case_insensitive_extension(self, parser):
        """Test that file extensions are case-insensitive."""
        content = b"Test content"
        
        # Should work with uppercase extension
        text1, _ = await parser.extract_text("file.TXT", content)
        text2, _ = await parser.extract_text("file.txt", content)
        
        assert text1 == text2
    
    @pytest.mark.asyncio
    async def test_multiple_dots_in_filename(self, parser):
        """Test filename with multiple dots."""
        content = b"Test content"
        text, pages = await parser.extract_text("my.file.name.txt", content)
        
        assert text == "Test content"
        assert pages == 1
    
    @pytest.mark.asyncio
    async def test_supported_formats_list(self, parser):
        """Test that error message lists supported formats."""
        try:
            await parser.extract_text("file.xyz", b"content")
        except UnsupportedFormatError as e:
            error_msg = str(e)
            assert "pdf" in error_msg.lower()
            assert "docx" in error_msg.lower()
            assert "txt" in error_msg.lower()


class TestBackwardCompatibility:
    """Test backward compatibility with legacy extract_text function."""
    
    @pytest.mark.asyncio
    async def test_legacy_function_works(self):
        """Test that legacy extract_text function still works."""
        from backend.document_parser import extract_text
        
        content = b"Legacy test content"
        text, pages = await extract_text("legacy.txt", content)
        
        assert text == "Legacy test content"
        assert pages == 1
